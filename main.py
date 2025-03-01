# import os
# import spacy
# import streamlit as st
# from pdfminer.high_level import extract_text as extract_text_from_pdf
# from docx import Document
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity

# st.title("📄 Resume Screening App")
# st.write("Upload your resume (PDF/DOCX) and enter a job description to get a match score.")

# # User inputs job description
# job_description = st.text_area("Enter Job Description:")

# # User uploads resume file (PDF/DOCX)
# uploaded_file = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])

# if uploaded_file and job_description:
#     # Extract text from PDF/DOCX
#     resume_text = (
#         extract_text_from_pdf(uploaded_file) if uploaded_file.name.endswith(".pdf") 
#         else "\n".join([para.text for para in Document(uploaded_file).paragraphs])
#     ).lower()

#     # Load NLP model and process text
#     nlp = spacy.load("en_core_web_sm")
    
#     # Extract keywords from resume and job description
#     resume_keywords = {token.text for token in nlp(resume_text) if token.is_alpha and not token.is_stop}
#     job_desc_keywords = {token.text for token in nlp(job_description) if token.is_alpha and not token.is_stop}

#     # Calculate match score using TF-IDF and cosine similarity
#     vectorizer = TfidfVectorizer()
#     tfidf_matrix = vectorizer.fit_transform([" ".join(job_desc_keywords), " ".join(resume_keywords)])
#     score = round(cosine_similarity(tfidf_matrix[0], tfidf_matrix[1])[0][0] * 100, 2)

#     # Find matched keywords
#     matched_keywords = resume_keywords.intersection(job_desc_keywords)

#     # Display results
#     st.subheader("Results:")
#     st.write(f"✅ **Match Score:** {score}%")
#     st.write(f"🔢 **Total Matches:** {len(matched_keywords)}")
    
#     # Expandable section for matched keywords
#     with st.expander("**View Matched Keywords**"):
#         st.write(", ".join(matched_keywords))


import os
import spacy
import streamlit as st
from pdfminer.high_level import extract_text as extract_text_from_pdf
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Ensure spaCy model is installed
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    os.system("python -m spacy download en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

st.title("📄 Resume Screening App")
st.write("Upload your resume (PDF/DOCX) and enter a job description to get a match score.")

# User inputs job description
job_description = st.text_area("Enter Job Description:")

# User uploads resume file (PDF/DOCX)
uploaded_file = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])

if uploaded_file and job_description:
    # Extract text from PDF/DOCX
    if uploaded_file.name.endswith(".pdf"):
        resume_text = extract_text_from_pdf(uploaded_file)
    else:
        doc = Document(uploaded_file)
        resume_text = "\n".join([para.text for para in doc.paragraphs])

    resume_text = resume_text.lower()

    # Extract keywords from resume and job description
    resume_keywords = {token.text for token in nlp(resume_text) if token.is_alpha and not token.is_stop}
    job_desc_keywords = {token.text for token in nlp(job_description) if token.is_alpha and not token.is_stop}

    # Calculate match score using TF-IDF and cosine similarity
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform([" ".join(job_desc_keywords), " ".join(resume_keywords)])
    score = round(cosine_similarity(tfidf_matrix[0], tfidf_matrix[1])[0][0] * 100, 2)

    # Find matched keywords
    matched_keywords = resume_keywords.intersection(job_desc_keywords)

    # Display results
    st.subheader("Results:")
    st.write(f"✅ **Match Score:** {score}%")
    st.write(f"🔢 **Total Matches:** {len(matched_keywords)}")
    
    # Expandable section for matched keywords
    with st.expander("**View Matched Keywords**"):
        st.write(", ".join(matched_keywords))
